"""Text extraction from various file formats.

Supports:
- PDF (via PyPDF2)
- DOCX (via python-docx)
- TXT, CSV, JSON, Markdown (plain read)
"""

import io
import json
import logging

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

# Supported MIME types mapped to extraction functions
SUPPORTED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
    "text/csv": "csv",
    "application/json": "json",
    "text/markdown": "md",
    "text/x-markdown": "md",
}


def extract_text(
    content: bytes, mime_type: str | None = None, filename: str | None = None
) -> str:
    """Extract text from file content based on MIME type or filename extension.

    Args:
        content: Raw file bytes.
        mime_type: MIME type of the file. If not provided, inferred from filename.
        filename: Original filename, used for extension-based type inference.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If file type cannot be determined or is unsupported.
        RuntimeError: If text extraction fails.
    """
    file_type = _resolve_file_type(mime_type, filename)

    extractors = {
        "pdf": _extract_pdf,
        "docx": _extract_docx,
        "txt": _extract_plain,
        "csv": _extract_plain,
        "json": _extract_json,
        "md": _extract_plain,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    try:
        text = extractor(content)
        logger.info(
            "Extracted %d characters from %s file",
            len(text),
            file_type,
        )
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {file_type} file: {e}") from e


def _resolve_file_type(mime_type: str | None, filename: str | None) -> str:
    """Resolve the file type from MIME type or filename extension.

    Args:
        mime_type: MIME type string.
        filename: Filename with extension.

    Returns:
        Normalized file type string (pdf, docx, txt, csv, json).

    Raises:
        ValueError: If type cannot be determined.
    """
    # Try MIME type first
    if mime_type:
        file_type = SUPPORTED_MIME_TYPES.get(mime_type)
        if file_type:
            return file_type

    # Fall back to filename extension
    if filename:
        ext = filename.rsplit(".", maxsplit=1)[-1].lower() if "." in filename else ""
        extension_map = {
            "pdf": "pdf",
            "docx": "docx",
            "txt": "txt",
            "csv": "csv",
            "json": "json",
            "text": "txt",
            "md": "md",
            "markdown": "md",
        }
        file_type = extension_map.get(ext)
        if file_type:
            return file_type

    raise ValueError(
        f"Cannot determine file type. MIME: {mime_type}, filename: {filename}. "
        f"Supported types: {list(SUPPORTED_MIME_TYPES.keys())}"
    )


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes.

    Args:
        content: Raw PDF file bytes.

    Returns:
        Concatenated text from all pages.
    """
    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            pages.append(page_text.strip())
            logger.debug("PDF page %d: extracted %d chars", i + 1, len(page_text))
        else:
            logger.warning(
                "PDF page %d: no text extracted (may be scanned/image)", i + 1
            )

    if not pages:
        raise RuntimeError(
            "No text could be extracted from PDF. "
            "The file may be scanned/image-based and require OCR."
        )

    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes.

    Extracts text from paragraphs and tables.

    Args:
        content: Raw DOCX file bytes.

    Returns:
        Concatenated text from all paragraphs and table cells.
    """
    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    if not parts:
        raise RuntimeError("No text could be extracted from DOCX file.")

    return "\n".join(parts)


def _extract_plain(content: bytes) -> str:
    """Extract text from plain text or CSV bytes.

    Tries UTF-8 first, falls back to latin-1.

    Args:
        content: Raw text file bytes.

    Returns:
        Decoded text string.
    """
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        logger.warning("UTF-8 decode failed, falling back to latin-1")
        return content.decode("latin-1")


def _extract_json(content: bytes) -> str:
    """Extract text from JSON bytes.

    For structured JSON, converts to a readable text representation.
    Handles both JSON objects and arrays.

    Args:
        content: Raw JSON file bytes.

    Returns:
        Pretty-printed JSON string or flattened text representation.
    """
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")

    # Validate it's actual JSON
    data = json.loads(text)

    # If it's a simple string value, return it directly
    if isinstance(data, str):
        return data

    # For structured data, return pretty-printed JSON
    # This preserves structure while making it readable for the LLM
    return json.dumps(data, indent=2, ensure_ascii=False)
